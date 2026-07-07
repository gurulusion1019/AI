"""
core/protocol_parser.py
Adaptive protocol document parser.
Accepts any text format. Extracts MeasureDefinitions via Claude API.
Identifies gaps and generates targeted clarification questions.
"""
import os
import re
import json
import time
import logging
from typing import Optional
import requests

from .models import (
    MeasureDefinition, AggregationType, FilterCondition,
    DataError, ErrorCategory, ErrorSeverity, ResolutionOption
)

logger = logging.getLogger(__name__)

ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
CLAUDE_MODEL = "claude-sonnet-4-5"


# ── Text extraction from various formats ───────────────────────────────────

def extract_text_from_docx(content: bytes) -> str:
    """Extract plain text from a .docx file."""
    try:
        import io
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)
        return '\n'.join(paragraphs)
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", exc)
        return ""


def extract_text_from_pdf(content: bytes) -> str:
    """Extract plain text from a PDF."""
    try:
        import io
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return '\n'.join(text_parts)
    except Exception as exc:
        logger.warning("PDF extraction failed: %s", exc)
        return ""


def extract_protocol_text(content: bytes, file_name: str) -> str:
    """Route to the right extractor based on file extension."""
    suffix = file_name.lower().split('.')[-1]
    if suffix == 'docx':
        return extract_text_from_docx(content)
    elif suffix == 'pdf':
        return extract_text_from_pdf(content)
    elif suffix in ('txt', 'md'):
        for enc in ('utf-8', 'utf-8-sig', 'latin-1'):
            try:
                return content.decode(enc)
            except UnicodeDecodeError:
                continue
    elif suffix in ('xlsx', 'xls'):
        try:
            import io
            import pandas as pd
            df = pd.read_excel(io.BytesIO(content))
            return df.to_string(index=False)
        except Exception as exc:
            logger.warning("Excel protocol extraction failed: %s", exc)
    return content.decode('utf-8', errors='replace')


# ── Claude API call ────────────────────────────────────────────────────────

def _call_claude(prompt: str, system: str = "", retries: int = 3) -> Optional[str]:
    if not ANTHROPIC_API_KEY:
        logger.warning("ANTHROPIC_API_KEY not set — using stub extraction")
        return None

    headers = {
        "x-api-key": ANTHROPIC_API_KEY,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    messages = [{"role": "user", "content": prompt}]
    payload: dict = {
        "model": CLAUDE_MODEL,
        "max_tokens": 16000,   # was 4096 — too small for ~30 verbose measures, truncated the JSON
        "messages": messages,
    }
    if system:
        payload["system"] = system

    # Stream the response. A long extraction (~30 measures ≈ 6-8k tokens) takes
    # well over 60s to produce as a single blob, which tripped the old read
    # timeout. Streaming delivers tokens continuously so the (connect, read)
    # timeout applies per-chunk, not to the whole generation.
    payload["stream"] = True

    for attempt in range(retries):
        try:
            with requests.post(
                "https://api.anthropic.com/v1/messages",
                headers=headers,
                json=payload,
                timeout=(15, 120),   # 15s to connect, 120s max gap between chunks
                stream=True,
            ) as resp:
                if resp.status_code == 529:          # overloaded — back off and retry
                    resp.close()
                    time.sleep(2 ** attempt)
                    continue
                if resp.status_code != 200:
                    logger.error("Claude API %s: %s", resp.status_code, resp.text[:300])
                    return None

                parts: list[str] = []
                for line in resp.iter_lines(decode_unicode=True):
                    if not line or not line.startswith("data:"):
                        continue
                    data = line[len("data:"):].strip()
                    if data == "[DONE]":
                        break
                    try:
                        evt = json.loads(data)
                    except json.JSONDecodeError:
                        continue
                    etype = evt.get("type")
                    if etype == "content_block_delta":
                        delta = evt.get("delta", {})
                        if delta.get("type") == "text_delta":
                            parts.append(delta.get("text", ""))
                    elif etype == "error":
                        logger.error("Claude stream error: %s", evt.get("error"))
                        return None

                text = "".join(parts)
                if text:
                    return text
                logger.error("Claude returned an empty stream")
                return None

        except requests.RequestException as exc:
            logger.error("Claude API request error (attempt %d/%d): %s",
                         attempt + 1, retries, exc)
            if attempt < retries - 1:
                time.sleep(2 ** attempt)
    return None


# ── Extraction ─────────────────────────────────────────────────────────────

EXTRACTION_SYSTEM = """You are a business intelligence analyst extracting metric definitions
from a protocol document. Extract ALL metrics described, even if described vaguely.
Flag low-confidence extractions rather than guessing. Preserve the source excerpt.
Return ONLY a valid JSON array. No markdown, no explanation."""

EXTRACTION_PROMPT_TEMPLATE = """Extract all metric/measure definitions from the following protocol document.

Available columns in the data (for base_column matching):
{available_columns}

For each metric, return:
{{
  "name": "snake_case_name",
  "display_name": "Human Readable Name",
  "description": "what it measures",
  "aggregation": "SUM|COUNT|COUNTD|AVG|MIN|MAX|RATIO|CUSTOM",
  "base_column": "column_name or null",
  "base_table": "table_name or null",
  "numerator": "column or expression for ratio (null if not ratio)",
  "denominator": "column or expression for ratio (null if not ratio)",
  "filter_conditions": [{{"column":"...", "operator":"=", "value":"..."}}],
  "time_intelligence": "YTD|MTD|ROLLING_12M|PRIOR_PERIOD|null",
  "dimensions": ["dim_table_name"],
  "format_string": "$#,##0 or 0.0% or #,##0",
  "is_kpi": true/false,
  "dependencies": ["other_measure_name"],
  "confidence": 0.0-1.0,
  "source_excerpt": "verbatim text this came from",
  "needs_clarification": true/false,
  "clarification_questions": ["question1", "question2"]
}}

PROTOCOL DOCUMENT:
{protocol_text}

Return ONLY the JSON array."""


def _parse_measure_json(raw: str) -> Optional[list]:
    """
    Parse Claude's measure-extraction response into a list of dicts.

    Robust to three real-world failure modes we hit:
      1. Response wrapped in ```json fences.
      2. Response is an object like {"measures": [...]} instead of a bare array.
      3. Response truncated mid-array (output hit max_tokens) -> salvage every
         COMPLETE object so we return 28 measures instead of falling to a stub.

    Returns a list on success, or None if nothing usable could be recovered.
    """
    cleaned = re.sub(r'```(?:json)?', '', raw).strip()

    # Fast path: clean parse.
    try:
        data = json.loads(cleaned)
        if isinstance(data, list):
            return data
        if isinstance(data, dict):
            return data.get('measures') or data.get('metrics') or [data]
    except json.JSONDecodeError:
        pass

    # Slice to the outermost array and retry.
    start = cleaned.find('[')
    end = cleaned.rfind(']')
    if start != -1 and end > start:
        try:
            data = json.loads(cleaned[start:end + 1])
            if isinstance(data, list):
                return data
        except json.JSONDecodeError:
            pass

    # Salvage a truncated array: walk it and keep every complete {...} object.
    if start != -1:
        objs, buf, depth = [], '', 0
        in_str = esc = False
        for ch in cleaned[start + 1:]:
            buf += ch
            if esc:
                esc = False
                continue
            if ch == '\\':
                esc = True
                continue
            if ch == '"':
                in_str = not in_str
            elif not in_str:
                if ch == '{':
                    depth += 1
                elif ch == '}':
                    depth -= 1
                    if depth == 0:
                        try:
                            objs.append(json.loads(buf.strip().lstrip(',').strip()))
                        except json.JSONDecodeError:
                            pass
                        buf = ''
        if objs:
            logger.warning("Salvaged %d complete measures from a truncated response", len(objs))
            return objs

    return None


def _stub_extraction(protocol_text: str, reason: str = "no API key") -> list[dict]:
    """
    Fallback when real extraction is unavailable.
    Returns a minimal stub so the pipeline can still be tested.
    `reason` is logged so we never again mistake a parse failure for a missing key.
    """
    logger.warning("Using stub extraction (%s) — real measures were NOT extracted", reason)
    return [
        {
            "name": "total_revenue",
            "display_name": "Total Revenue",
            "description": "Sum of all sale amounts",
            "aggregation": "SUM",
            "base_column": "sale_amount",
            "base_table": None,
            "numerator": None,
            "denominator": None,
            "filter_conditions": [],
            "time_intelligence": None,
            "dimensions": [],
            "format_string": "$#,##0",
            "is_kpi": True,
            "dependencies": [],
            "confidence": 0.95,
            "source_excerpt": "(stub — no API key)",
            "needs_clarification": False,
            "clarification_questions": [],
        }
    ]


def extract_measures(
    protocol_text: str,
    available_columns: list[str],
) -> list[MeasureDefinition]:
    """
    Parse the protocol document and return a list of MeasureDefinitions.
    Uses Claude API. Falls back to stub if no key.
    """
    col_str = ', '.join(available_columns[:80]) if available_columns else '(unknown)'
    prompt = EXTRACTION_PROMPT_TEMPLATE.format(
        available_columns=col_str,
        protocol_text=protocol_text[:12000],   # stay within context
    )

    raw = _call_claude(prompt, system=EXTRACTION_SYSTEM)

    items = None
    if raw:
        items = _parse_measure_json(raw)
        if items is None:
            logger.error(
                "Measure JSON unparseable even after salvage. raw_len=%d tail=%r",
                len(raw), raw[-300:],
            )

    if not items:
        reason = "no API key / empty response" if not raw else "unparseable API response"
        items = _stub_extraction(protocol_text, reason=reason)
    else:
        logger.info("Extracted %d measures from protocol", len(items))

    measures = []
    for item in items:
        # Parse filter conditions
        filters = []
        for fc in item.get('filter_conditions') or []:
            try:
                filters.append(FilterCondition(
                    column=fc.get('column', ''),
                    operator=fc.get('operator', '='),
                    value=fc.get('value', ''),
                ))
            except Exception:
                pass

        # Map aggregation string to enum
        agg_str = (item.get('aggregation') or 'SUM').upper()
        try:
            agg = AggregationType(agg_str)
        except ValueError:
            agg = AggregationType.CUSTOM

        m = MeasureDefinition(
            name=item.get('name', 'unnamed_measure'),
            display_name=item.get('display_name', item.get('name', '')),
            description=item.get('description', ''),
            aggregation=agg,
            numerator=item.get('numerator'),
            denominator=item.get('denominator'),
            base_column=item.get('base_column'),
            base_table=item.get('base_table'),
            filter_conditions=filters,
            time_intelligence=item.get('time_intelligence'),
            dimensions=item.get('dimensions') or [],
            format_string=item.get('format_string', '#,##0'),
            is_kpi=bool(item.get('is_kpi', False)),
            dependencies=item.get('dependencies') or [],
            confidence=float(item.get('confidence', 0.8)),
            source_excerpt=item.get('source_excerpt', ''),
            needs_clarification=bool(item.get('needs_clarification', False)),
            clarification_questions=item.get('clarification_questions') or [],
        )
        measures.append(m)

    return measures


# ── Conflict detection ─────────────────────────────────────────────────────

def detect_protocol_conflicts(
    protocol_text: str,
    measures: list[MeasureDefinition],
) -> list[DataError]:
    """
    Detect conflicting definitions for the same measure name.
    Ask Claude to find contradictions.
    """
    if not ANTHROPIC_API_KEY:
        return []

    errors = []

    # Group measures by similar names
    from collections import defaultdict
    name_groups: dict[str, list] = defaultdict(list)
    for m in measures:
        key = re.sub(r'[^a-z]', '', m.name.lower())
        name_groups[key].append(m)

    for key, group in name_groups.items():
        if len(group) < 2:
            continue

        # Two measures with same base name — check if formulas differ
        formulas = [m.source_excerpt for m in group if m.source_excerpt]
        if len(set(formulas)) < 2:
            continue

        e = DataError(
            category=ErrorCategory.PROTOCOL,
            sub_type="CONFLICTING_DEFINITIONS",
            severity=ErrorSeverity.HIGH,
            affected_table="protocol",
            affected_column=group[0].display_name,
            detail=(
                f"'{group[0].display_name}' is defined {len(group)} times "
                f"with different formulas. Using the wrong one will produce incorrect results."
            ),
            sample_data=[
                {"version": f"Definition {i+1}", "excerpt": m.source_excerpt[:200]}
                for i, m in enumerate(group)
            ],
            options=[
                ResolutionOption(str(chr(97 + i)), f"Use definition {i+1}",
                    m.source_excerpt[:120],
                    "Selected formula applied", 0, 0.0, True, i == 0)
                for i, m in enumerate(group)
            ] + [
                ResolutionOption(str(chr(97 + len(group))),
                    "Build both as separate measures",
                    "Create one measure per definition for comparison.",
                    "Both measures created", 0, 0.0, True, False),
                ResolutionOption(str(chr(97 + len(group) + 1)),
                    "Skip this measure",
                    "Exclude from model. Re-add with clarified protocol.",
                    "Measure skipped", 0, 0.0, False, False),
            ],
            recommended_option_id='a',
        )
        errors.append(e)

    return errors


# ── Dependency cycle detection ─────────────────────────────────────────────

def detect_dependency_cycles(measures: list[MeasureDefinition]) -> list[DataError]:
    """Topological sort to find circular measure dependencies."""
    name_map = {m.name: m for m in measures}
    visited: set[str] = set()
    rec_stack: set[str] = set()
    cycles_found: list[list[str]] = []

    def dfs(name: str, path: list[str]):
        visited.add(name)
        rec_stack.add(name)
        path = path + [name]
        m = name_map.get(name)
        if m:
            for dep in m.dependencies:
                if dep not in visited:
                    dfs(dep, path)
                elif dep in rec_stack:
                    idx = path.index(dep) if dep in path else 0
                    cycles_found.append(path[idx:] + [dep])
        rec_stack.discard(name)

    for m in measures:
        if m.name not in visited:
            dfs(m.name, [])

    errors = []
    for cycle in cycles_found:
        e = DataError(
            category=ErrorCategory.PROTOCOL,
            sub_type="DEPENDENCY_CYCLE",
            severity=ErrorSeverity.BLOCKER,
            affected_table="protocol",
            affected_column=" → ".join(cycle),
            detail=(
                f"Circular dependency detected: {' → '.join(cycle)}. "
                f"DAX cannot resolve circular measure references."
            ),
            options=[
                ResolutionOption('a', f"Break link: remove '{cycle[-2]}' dependency",
                    "Remove the last dependency in the cycle.",
                    "Cycle resolved, measure redefined", 0, 0.0, True, True),
                ResolutionOption('b', "Manually redefine one measure",
                    "Clarify the independent base calculation for one of the measures.",
                    "Cycle resolved via clarification", 0, 0.0, True, False),
            ],
            recommended_option_id='a',
        )
        errors.append(e)

    return errors


# ── DAX generation ─────────────────────────────────────────────────────────

DAX_SYSTEM = """You are a Power BI DAX expert. Generate correct, validated DAX measure formulas.
Return ONLY a JSON object with field "dax". No explanation, no markdown."""

def _remap_dax_tables(dax: str, table_schema: dict) -> str:
    """
    Rewrite table references in a DAX formula to match ACTUAL table names.

    Claude may generate DAX using table names from the protocol document
    (e.g. "FACT_Orders") that don't match the real table names derived from
    filenames (e.g. "retail_fact_orders"). This remaps them by:
    1. Exact match (already correct) → leave alone
    2. Fuzzy match on table name → remap
    3. Column-based match (which real table has this column) → remap

    table_schema: {actual_table_name: [column names]}
    """
    import re as _re

    actual_tables = list(table_schema.keys())
    if not actual_tables:
        return dax

    # Build a lowercase lookup for fuzzy table matching
    def _norm(s):
        return _re.sub(r'[^a-z0-9]', '', s.lower())

    actual_norm = {_norm(t): t for t in actual_tables}

    # Build column → table index (for column-based resolution)
    col_to_tables = {}
    for t, cols in table_schema.items():
        for c in cols:
            col_to_tables.setdefault(c.lower(), []).append(t)

    # Find all 'TableName'[Column] or TableName[Column] references
    # Pattern captures optional quotes, table name, and column
    pattern = r"'?([A-Za-z_][A-Za-z0-9_ ]*?)'?\[([^\]]+)\]"

    def _replace(match):
        ref_table = match.group(1).strip()
        col = match.group(2).strip()

        # 1. Already an actual table? keep it
        if ref_table in actual_tables:
            return f"'{ref_table}'[{col}]"

        # 2. Fuzzy match on normalized table name
        n = _norm(ref_table)
        if n in actual_norm:
            return f"'{actual_norm[n]}'[{col}]"
        # partial fuzzy (contains)
        for an, at in actual_norm.items():
            if n and (n in an or an in n):
                return f"'{at}'[{col}]"

        # 3. Column-based — which real table has this column?
        candidates = col_to_tables.get(col.lower(), [])
        if len(candidates) == 1:
            return f"'{candidates[0]}'[{col}]"
        elif len(candidates) > 1:
            # Prefer a fact-like table if ambiguous
            for cand in candidates:
                if 'fact' in cand.lower():
                    return f"'{cand}'[{col}]"
            return f"'{candidates[0]}'[{col}]"

        # 4. No resolution — leave as-is (will error visibly, better than silent wrong)
        return match.group(0)

    return _re.sub(pattern, _replace, dax)


def generate_dax(
    measure: MeasureDefinition,
    table_schema: dict[str, list[str]],
    all_measures: list[str],
) -> str:
    """
    Generate a DAX formula for a measure.
    Falls back to a template if no API key.
    """
    if not ANTHROPIC_API_KEY:
        return _template_dax(measure)

    schema_str = json.dumps(table_schema, indent=2)
    filter_str = json.dumps([
        {"column": f.column, "operator": f.operator, "value": f.value}
        for f in measure.filter_conditions
    ])

    prompt = f"""Generate a DAX measure formula for Power BI.

Measure: {measure.display_name}
Description: {measure.description}
Aggregation: {measure.aggregation.value}
Base column: {measure.base_column} (in table: {measure.base_table or 'any'})
Numerator: {measure.numerator or 'n/a'}
Denominator: {measure.denominator or 'n/a'}
Filter conditions: {filter_str}
Time intelligence: {measure.time_intelligence or 'none'}
Format: {measure.format_string}
Dependencies on other measures: {measure.dependencies}
Existing measure names: {all_measures[:20]}

Table schema:
{schema_str}

Return JSON: {{"dax": "MEASURE_NAME = <formula here>"}}"""

    raw = _call_claude(prompt, system=DAX_SYSTEM)
    if raw:
        try:
            raw = re.sub(r'```(?:json)?', '', raw).strip()
            result = json.loads(raw)
            dax = result.get("dax", _template_dax(measure))
            return _remap_dax_tables(dax, table_schema)
        except Exception:
            pass

    return _remap_dax_tables(_template_dax(measure), table_schema)


def _template_dax(m: MeasureDefinition) -> str:
    """Generate a template DAX formula without the API."""
    col = f"'{m.base_table}'[{m.base_column}]" if m.base_table and m.base_column else "[value]"
    name = m.display_name or m.name

    templates = {
        AggregationType.SUM:    f"{name} = SUM({col})",
        AggregationType.COUNT:  f"{name} = COUNT({col})",
        AggregationType.COUNTD: f"{name} = DISTINCTCOUNT({col})",
        AggregationType.AVG:    f"{name} = AVERAGE({col})",
        AggregationType.MIN:    f"{name} = MIN({col})",
        AggregationType.MAX:    f"{name} = MAX({col})",
        AggregationType.RATIO:  (
            f"{name} = DIVIDE(SUM('{m.base_table}'[{m.numerator}]), "
            f"SUM('{m.base_table}'[{m.denominator}]), 0)"
            if m.numerator and m.denominator else f"{name} = DIVIDE([numerator], [denominator], 0)"
        ),
        AggregationType.CUSTOM: f"{name} = -- TODO: define custom formula\n    SUM({col})",
    }

    base = templates.get(m.aggregation, f"{name} = SUM({col})")

    # Wrap with time intelligence
    if m.time_intelligence == "YTD":
        base = f"{name} = CALCULATE({base.split('=',1)[1].strip()}, DATESYTD('dim_date'[date]))"
    elif m.time_intelligence == "MTD":
        base = f"{name} = CALCULATE({base.split('=',1)[1].strip()}, DATESMTD('dim_date'[date]))"

    # Wrap with filters
    if m.filter_conditions:
        filters = ", ".join(
            f"'{f.column}'[{f.column}] = \"{f.value}\""
            for f in m.filter_conditions[:3]
        )
        inner = base.split('=', 1)[1].strip()
        base = f"{name} = CALCULATE({inner}, {filters})"

    return base